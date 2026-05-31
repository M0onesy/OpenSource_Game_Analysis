#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_paper.py — 由 python-docx 生成《大模型公司开源 vs 闭源战略博弈分析》Word 文档。

按作业要求组织（第 1 页含班级/姓名/学号/各人贡献；正文：Game 介绍 / 要素识别 /
最优反应与均衡 / 改进一 / 改进二 / 改进三 / 结论），不含摘要与参考文献。

特性：US Letter、窄边距（0.5 英寸）、页脚页码；全文中文字体（Noto Sans CJK SC，含 eastAsia）；
正文为连续学术散文（不使用高亮框、不使用项目符号）；内嵌 6 张中文图，图号按出现顺序为 图1…图6。
支付不再外生赋值：静态博弈支付由 §2.4 支付函数实例化计算，重复博弈单期支付由差异化伯特兰模型推导。

运行：cd 代码/ && python build_paper.py   输出：../论文/大模型博弈论分析_论文.docx
"""

import os
from _cjk_support import configure_console_utf8, pick_cjk_font
from _png_support import strip_incorrect_iccp_chunk
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

configure_console_utf8()

HERE = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.normpath(os.path.join(HERE, "..", "图表"))
OUT_DIR = os.path.normpath(os.path.join(HERE, "..", "论文"))
os.makedirs(OUT_DIR, exist_ok=True)
OUT_DOCX = os.path.join(OUT_DIR, "大模型博弈论分析_论文.docx")

CJK = pick_cjk_font()
INK = RGBColor(0x1A, 0x1A, 0x1A)
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0xC0, 0x50, 0x4D)
GREEN = RGBColor(0x2F, 0x6B, 0x3F)
HDR_FILL = "1F4E79"
SUB_FILL = "D5E1EE"
ZEBRA_FILL = "EEF3F8"
CALLOUT_FILL = "FBF1E6"   # 浅暖色高亮框

# 图号 = 出现顺序（1..6）→（文件名, 宽度英寸）；窄边距内容宽约 7.5"
FIGS = {
    1: ("fig1_overview.png", 6.6),
    2: ("fig2_payoff_matrix.png", 4.3),
    3: ("fig3_game_tree.png", 6.9),
    4: ("fig4_evolutionary.png", 7.0),
    5: ("fig5_bayesian_signal.png", 7.0),
    6: ("fig6_repeated_game.png", 7.0),
}


# ----------------------------------------------------------------- helpers
def set_cjk_on_style(style, font_name=CJK):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), font_name)


def _set_run_cjk(run, font_name=CJK):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), font_name)


def add_para(doc, runs, *, align=None, space_before=0, space_after=4,
             line=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line is not None:
        pf.line_spacing = line
    if isinstance(runs, str):
        runs = [(runs, {})]
    for text, opt in runs:
        r = p.add_run(text)
        _set_run_cjk(r)
        r.font.size = Pt(opt.get("size", 10.5))
        r.font.bold = opt.get("bold", False)
        r.font.italic = opt.get("italic", False)
        r.font.color.rgb = opt.get("color", INK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    _set_run_cjk(r)
    if level == 1:
        r.font.size = Pt(15)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    else:
        r.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
    r.font.color.rgb = PRIMARY
    r.font.bold = True
    return p


def _shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def _set_cell_margins(cell, top=40, bottom=40, left=90, right=90):
    tcpr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("left", left), ("bottom", bottom),
                     ("right", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcpr.append(m)


def _fill_cell(cell, text, *, bold=False, size=9, color=INK, align=None,
               fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    _set_run_cjk(r)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
    if fill:
        _shade_cell(cell, fill)
    _set_cell_margins(cell)


def add_table(doc, header, rows, *, widths=None, header_fill=HDR_FILL,
              zebra=True, size=9, caption=None):
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    total = 6.5
    if widths is None:
        widths = [total / ncol] * ncol
    hdr = table.rows[0].cells
    for j, htext in enumerate(header):
        _fill_cell(hdr[j], htext, bold=True, size=size,
                   color=RGBColor(0xFF, 0xFF, 0xFF),
                   align=WD_ALIGN_PARAGRAPH.CENTER, fill=header_fill)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        rfill = ZEBRA_FILL if (zebra and i % 2 == 1) else None
        for j, val in enumerate(row):
            al = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            _fill_cell(cells[j], val, size=size, align=al, fill=rfill)
    for j, w in enumerate(widths):
        for row in table.rows:
            row.cells[j].width = Inches(w)
    if caption:
        add_para(doc, [(caption, {"size": 9, "italic": True,
                                  "color": RGBColor(0x55, 0x55, 0x55)})],
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=5)
    return table


def add_callout(doc, label, text):
    """单行单元格高亮框，用于突出每节核心结论。左侧粗边 + 暖色底。"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Inches(6.5)
    tcpr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")

    def _bd(side, sz, color):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), color)
        return e
    # 子元素顺序必须为 top, left, bottom, right
    borders.append(_bd("top", 4, "E6C9A8"))
    borders.append(_bd("left", 30, "C0504D"))
    borders.append(_bd("bottom", 4, "E6C9A8"))
    borders.append(_bd("right", 4, "E6C9A8"))
    tcpr.append(borders)                      # tcBorders 必须在 shd / tcMar 之前
    _shade_cell(cell, CALLOUT_FILL)
    _set_cell_margins(cell, top=55, bottom=55, left=150, right=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(label + "　")
    _set_run_cjk(r1); r1.font.size = Pt(10); r1.font.bold = True; r1.font.color.rgb = ACCENT
    r2 = p.add_run(text)
    _set_run_cjk(r2); r2.font.size = Pt(10); r2.font.color.rgb = INK
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
    return t


def add_figure(doc, fignum, caption):
    fname, width = FIGS[fignum]
    path = os.path.join(FIG_DIR, fname)
    strip_incorrect_iccp_chunk(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    add_para(doc, [(caption, {"size": 9, "italic": True,
                              "color": RGBColor(0x55, 0x55, 0x55)})],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _grey(run):
        _set_run_cjk(run); run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    def _field(instr):
        fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), instr)
        r = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:eastAsia"):
            rf.set(qn(a), CJK)
        rpr.append(rf)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18"); rpr.append(sz)
        col = OxmlElement("w:color"); col.set(qn("w:val"), "808080"); rpr.append(col)
        r.append(rpr)
        t = OxmlElement("w:t"); t.text = "1"; r.append(t); fld.append(r)
        p._p.append(fld)

    _grey(p.add_run("第 ")); _field("PAGE")
    _grey(p.add_run(" 页  /  共 ")); _field("NUMPAGES"); _grey(p.add_run(" 页"))


def add_rule(doc, color="1F4E79", size=6, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    ppr.insert(0, pbdr)
    return p


# ----------------------------------------------------------------- build
def build():
    doc = Document()
    set_cjk_on_style(doc.styles["Normal"])
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"].font.color.rgb = INK
    for sname in ("Heading 1", "Heading 2", "Heading 3", "Title"):
        try:
            set_cjk_on_style(doc.styles[sname])
        except KeyError:
            pass

    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(0.5))          # 窄边距
    add_page_number_footer(sec)

    BODY = 10.5     # 正文字号
    LINE = 1.10     # 行距

    # ========================================================= 封面
    sp = doc.add_paragraph(); sp.paragraph_format.space_before = Pt(4)
    add_para(doc, [("博弈论课程大作业", {"size": 13, "bold": True,
                    "color": RGBColor(0x80, 0x80, 0x80)})],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_para(doc, [("大模型公司「开源 vs 闭源」", {"size": 24, "bold": True, "color": PRIMARY})],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, [("战略博弈分析", {"size": 24, "bold": True, "color": PRIMARY})],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    add_para(doc, [("基于 2023–2026 年全球大模型产业竞争的多层博弈研究",
                    {"size": 12, "italic": True, "color": RGBColor(0x55, 0x55, 0x55)})],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    add_rule(doc, space_after=10)

    info = [
        ("课程名称", "博弈论"),
        ("作业类型", "课程大作业（小组）"),
        ("完成时间", "2026 年 5 月"),
        ("数据基线", "2025 年 1 月（DeepSeek-R1 冲击）— 2026 年 5 月 28 日"),
        ("方法特征", "支付由支付函数与产业定价模型推导，全部图表由代码 1:1 复现"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for k, v in info:
        cells = t.add_row().cells
        _fill_cell(cells[0], k, bold=True, size=10.5, fill=SUB_FILL,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _fill_cell(cells[1], v, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in t.rows:
        row.cells[0].width = Inches(1.9)
        row.cells[1].width = Inches(5.4)

    sp2 = doc.add_paragraph(); sp2.paragraph_format.space_after = Pt(7)
    add_para(doc, [("小组信息与成员分工", {"size": 13, "bold": True, "color": PRIMARY})],
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)
    add_para(doc, [("班级：______________________      组号：__________", {"size": 11})],
             space_after=8)
    add_table(doc,
        ["姓名", "学号", "主要分工与贡献"],
        [["______", "________________", "（如：建模与求解 / 代码与图表 / 数据核实 / 撰写）"],
         ["______", "________________", "________________________________________________"],
         ["______", "________________", "________________________________________________"],
         ["______", "________________", "________________________________________________"],
         ["______", "________________", "________________________________________________"]],
        widths=[1.2, 2.0, 4.0], size=9.5, zebra=False)
    sp3 = doc.add_paragraph(); sp3.paragraph_format.space_after = Pt(4)
    add_para(doc, [("提示：本封面为占位模板，请各小组据实填写班级、组号、姓名、学号与分工贡献。",
                    {"size": 9, "italic": True, "color": RGBColor(0x80, 0x80, 0x80)})],
             space_after=0)
    doc.add_page_break()

    # ========================================================= 一、Game 介绍
    add_heading(doc, "一、Game 介绍", 1)
    add_heading(doc, "1.1 产业背景：三轮转折", 2)
    add_para(doc, [("第一阶段（2023，闭源主导）。", {"bold": True, "size": BODY}),
                   ("OpenAI 发布 GPT-4，确立闭源 API 模式，输入价高达每百万 token 约 30 美元，与 Anthropic、Google "
                    "共同形成「闭源 + 高价 + 隐性价格协调」的先发护城河；同年 Meta 发布 Llama 2 开放权重，揭开开源序幕。"
                    "这一阶段的市场叙事是「能力来自规模、规模来自资本」，高昂的训练投入被视为质量的天然背书。",
                    {"size": BODY})], line=LINE, space_after=3)
    add_para(doc, [("第二阶段（2024，开源破局）。", {"bold": True, "size": BODY}),
                   ("Meta 持续发布 Llama 3 系列（Llama 3.1-405B 训练算力约 1.7 亿美元，Stanford AI Index 2025），阿里 Qwen "
                    "以 Apache 2.0 全面开源；12 月 26 日 DeepSeek 在 V3 技术报告（arXiv:2412.19437）中披露仅约 557.6 万美元"
                    "计算成本即完成训练，引发全球震动。开源阵营第一次在「能力」而非仅「可得性」上逼近闭源前沿。",
                    {"size": BODY})], line=LINE, space_after=3)
    add_para(doc, [("第三阶段（2025–2026，成本革命与监管分化）。", {"bold": True, "size": BODY}),
                   ("2025 年 1 月 20 日 DeepSeek-R1 发布，1 月 27 日 NVIDIA 单日下跌约 17%、市值蒸发约 5890 亿美元；"
                    "OpenAI 于 8 月 5 日发布 gpt-oss-120b/20b（自 GPT-2 以来首个开放权重模型），8 月 7 日推出 GPT-5，"
                    "输入价降至每百万 token 约 1.25 美元（较 GPT-4 下降约 95.8%）；Anthropic 旗舰价稳定在约 5/25 美元。"
                    "监管层面，欧盟《人工智能法案》的通用人工智能（GPAI）义务自 2025 年 8 月 2 日适用，加州 SB 53 于 2026 年 "
                    "1 月 1 日生效，美国第 14179 号行政令转向去管制，中国《人工智能+》行动（2025 年 8 月 27 日）鼓励开源。"
                    "进入 2026 年，OpenAI 发布 GPT-5.5，Anthropic 于 5 月 28 日发布 Claude Opus 4.8 并以约 9650 亿美元估值"
                    "反超 OpenAI（约 8520 亿美元）；DeepSeek 4 月 24 日发布 V4（MIT、100 万上下文），阿里 2 月发布 Qwen 3.5 并持续开源。",
                    {"size": BODY})], line=LINE, space_after=3)
    add_para(doc, [("两条印证性事实（Stanford AI Index 2025）：", {"bold": True, "size": 10}),
                   ("其一，GPT-3.5 等效推理价格在 18 个月内下降约 280 倍；其二，最强开源与最强闭源模型在主流基准上的能力差距"
                    "收敛至约 1.7 个百分点（早期流传的「17.5 个百分点」为口径误读，本文以约 1.7 个百分点为准）。这两条事实分别从"
                    "「价格」与「能力」两个维度，构成下文重复博弈与信号博弈的经验基础。", {"size": 10})],
             line=LINE, space_after=4)

    add_heading(doc, "1.2 核心博弈问题", 2)
    add_para(doc,
        "本文聚焦一个看似矛盾的现象：为何 Meta、DeepSeek、阿里坚持开源，而 OpenAI、Anthropic 坚持闭源？这种分化究竟是"
        "暂时的策略试探，还是稳定的博弈均衡？产业的演化路径能否被博弈论事前预测？要回答这些问题，需要把行业看作一个"
        "多参与人、多阶段、不完全信息的复杂博弈：开源者以免费权重换取生态绑定与人才品牌，闭源者以 API 与产品垄断换取"
        "直接收入与安全护城河，监管者通过合规成本改变各方支付，用户的迁移成本则决定锁定效应的强度。", line=LINE, space_after=3)
    add_para(doc,
        "本文的核心论点是：「开源—闭源分层共存」并非偶然，而是一个可由博弈论刻画、并能产生可证伪预测的稳定均衡现象。"
        "为论证这一点，下文依次用静态博弈、序贯博弈、贝叶斯信号博弈、无限期重复博弈与演化博弈五类工具逐层求解，"
        "并在每一步把模型结论对照 2023–2026 年的真实事件加以检验。需要强调的是，本文的支付不再外生设定：静态博弈的"
        "支付矩阵由第二节给出的支付函数实例化计算，重复博弈的单期支付由差异化伯特兰定价模型推导，从而使「均衡」结论"
        "建立在可追溯的原始量之上，而非循环论证。总体框架见图 1。", line=LINE, space_after=4)
    add_figure(doc, 1, "图 1：本研究博弈模型架构总览（五类参与人 + 基础静态/序贯博弈 + 三项扩展改进）。由 06_overview.py 生成。")

    # ========================================================= 二、要素识别
    add_heading(doc, "二、Game 要素识别", 1)
    add_heading(doc, "2.1 参与人", 2)
    add_table(doc,
        ["参与人", "类型", "旗舰（2026.5）", "估值 / 营收（2026）"],
        [["OpenAI", "闭源·高成本·先动领先者", "GPT-5 / 5.5", "估值约 8520 亿美元；消费端约 8 亿周活领先"],
         ["Anthropic", "闭源·高成本·企业端领先者", "Claude Opus 4.8 / 4.7", "估值约 9650 亿美元（5.28 反超）；营收年化约 470 亿美元"],
         ["Meta", "开源·高成本·生态型", "Llama 4（Behemoth 训练中）", "公司市值超 1.5 万亿美元；Llama 为成本中心"],
         ["DeepSeek", "开源·低成本·颠覆者", "V4（MIT，100 万上下文）", "私营；V4-Flash 输入约 0.14 美元/百万 token"],
         ["阿里通义", "开源·中成本·平台型", "Qwen 3.5 / Qwen3-VL", "HF 累计下载 2026.1 破 7 亿、2026.3 近 10 亿"]],
        widths=[1.0, 1.9, 1.6, 2.7], size=8.5,
        caption="表 1：五方参与人定位（2026 年 5 月）。OpenAI 消费端领先，Anthropic 企业 API/编码场景领先并于 5 月反超估值。")
    add_para(doc,
        "五方的策略分化并非偶然，而是由各自的成本结构、收入模式与生态位共同决定的。OpenAI 与 Anthropic 拥有最强的"
        "闭源旗舰与直接 API 收入，因而有动力维持闭源以保护定价权；Meta 把开源 Llama 作为「互补品商品化」战略——通过免费"
        "权重削弱竞争对手的 API 收入，同时反哺自家社交与广告主业，使模型本身从「利润中心」转为「需求放大器」；DeepSeek "
        "凭借极低的训练成本，开源几乎不损失直接收入，却能迅速积累全球声誉与人才吸引力；阿里则借开源把开发者导入阿里云"
        "生态，用算力与云服务变现。除主博弈五方外，Google（Gemini 闭源 + Gemma 开源）、xAI（Grok 双轨）、Mistral（逐步 "
        "API 化）等次要参与人虽不纳入核心求解，但通过改变市场的价格与生态环境，间接影响五方支付，可视为博弈的「环境变量」。",
        line=LINE, space_after=4)

    add_heading(doc, "2.2 行动空间", 2)
    add_para(doc, [("每个参与人的行动是二维向量 ", {"size": BODY}),
                   ("aᵢ = (sᵢ, pᵢ)", {"bold": True, "size": BODY, "color": PRIMARY}),
                   ("：战略维度 sᵢ ∈ {O 完全开源, H 半开源/开放权重, C 完全闭源}，定价维度 pᵢ ∈ {低, 中, 高}，"
                    "故完整纯策略共 9 个。为使核心均衡分析清晰，本文在主博弈中把 {O, H} 合并为「O」（非闭源），"
                    "并把定价维度通过支付函数内化——定价的策略互动单独放在第六节的重复博弈中处理。这样的降维不损失关键结论："
                    "开源与半开源在「放弃 API 直接收入、换取生态与品牌」这一根本权衡上是同向的，二者的差异主要体现为程度而非方向。",
                    {"size": BODY})], line=LINE, space_after=4)

    add_heading(doc, "2.3 行动顺序与信息结构", 2)
    add_para(doc,
        "行动顺序呈 Stackelberg 多领导者结构：t₁（2023.3）OpenAI 率先以 GPT-4 设定闭源高价基准；t₂（2023.7–2025.4）"
        "Meta 与 Anthropic 跟进；t₃（2024.12–2025.1）DeepSeek 与阿里再次调整。这恰好对应「基准设定者 → 开源响应者 → 颠覆者」"
        "的三波次历史。信息结构为不完全且不完美：私有信息包含真实训练成本 cᵢ、模型能力 θᵢ 与算力存量 Kᵢ；公共信号包含"
        "模型发布、基准跑分（MMLU、SWE-Bench、AIME 等）、API 价格、下载量以及技术报告的披露程度。DeepSeek 的颠覆之所以剧烈，"
        "正源于其「低成本类型」这一私有信息经 V3 技术报告突然披露，打破了原有的贝叶斯均衡（详见第五节）。", line=LINE, space_after=4)

    add_heading(doc, "2.4 支付函数：从原始量到效用", 2)
    add_para(doc,
        "本文不直接为博弈矩阵填入数字，而是先写出一个把各参与人选择映射为效用的支付函数，再由可追溯的原始量实例化它。"
        "对参与人 i，在自身策略 sᵢ 与对手策略组合 s₋ᵢ 下，其支付为：", line=LINE, space_after=2)
    add_para(doc, [("Uᵢ(sᵢ, s₋ᵢ) = Rᵢᴬᴾᴵ·𝟙[sᵢ=C] + αᵢ·E(n_O)·𝟙[sᵢ=O] + Vᵢ(sᵢ) − κ·(cᵢ/c_ref)·m(sᵢ)",
                    {"size": 11, "bold": True, "color": PRIMARY})],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=3)
    add_para(doc,
        "四个分量的经济含义与标定依据如下。第一项是 API 货币化收入，仅在闭源时为正：Rᵢᴬᴾᴵ = baseᵢ / n_C，其中 n_C 为"
        "闭源者人数，baseᵢ 度量该公司在闭源时可获取的高端 API 利润潜力。以人数相除刻画了竞争稀释——独家闭源者可独占高端"
        "市场，闭源者越多则份额与价格越被压低。baseᵢ 的序数排序有明确证据支撑：OpenAI 拥有约 8 亿周活与最强先发分发，"
        "故 base 最高；Meta 具备分发但以广告为主业，居中；DeepSeek 作为后入场者缺乏企业级分发，闭源变现能力最弱。"
        "第二项是开源生态价值，仅在开源时为正：E(n_O) = max(0, E₀ − τ·(n_O − 1))，随开源者人数 n_O 上升而拥挤折损"
        "（开发者、衍生模型与心智在 Llama、Qwen、DeepSeek 之间被分摊）；权重 αᵢ 度量开源生态对该公司主营的战略价值，"
        "Meta（广告与基建协同）与 DeepSeek（采用与人才）较高，OpenAI 因直接变现而较低。第三项 Vᵢ(sᵢ) 是品牌与战略价值，"
        "分开源品牌 V(O) 与闭源品牌 V(C)：DeepSeek 的开源品牌被「极低成本逼近前沿」的颠覆叙事显著放大，OpenAI 则在"
        "闭源前沿领导者形象上拥有最高的闭源品牌。第四项是成本负担：训练成本比 cᵢ/c_ref 直接取自真实数据，权重 κ 把它缩放"
        "到与其它效用项可比，维护乘子 m(sᵢ) 满足 m(C) > m(O)——闭源需独力承担前沿研发与推理基建，开源可由社区分担约四成。",
        line=LINE, space_after=3)
    add_para(doc,
        "本文唯一直接取自硬数据的原始量是训练算力成本：GPT-4 约 7800 万美元（含研发投入「超过 1 亿美元」）、Llama 3.1-405B "
        "约 1.7 亿美元、DeepSeek-V3 约 557.6 万美元，三者之比约为 14 : 30 : 1。其余参数（baseᵢ、αᵢ、Vᵢ 与共享参数 "
        "E₀、τ、κ、m）为按上述公开证据标定的序数效用，其「相对排序」是稳健结论，绝对数值仅用于使均衡可计算。"
        "完整参数取值见表 2。", line=LINE, space_after=4)
    add_table(doc,
        ["参与人", "训练成本(百万美元)", "成本比 cᵢ/c_ref", "API 强度 base", "生态权重 α", "开源品牌 V(O)", "闭源品牌 V(C)"],
        [["OpenAI", "78", "14.0", "12", "0.5", "1.0", "2.5"],
         ["Meta", "170", "30.5", "7", "1.4", "2.0", "0.5"],
         ["DeepSeek", "5.576", "1.0", "4", "1.2", "3.5", "0.3"]],
        widths=[0.95, 1.45, 1.15, 1.05, 0.95, 0.9, 0.85], size=8.5,
        caption="表 2：支付函数的参数标定。共享参数：独家开源生态 E₀=6、拥挤折损 τ=1.5、成本权重 κ=0.10、维护乘子 m(C)=1.0 / m(O)=0.6。")

    # ========================================================= 三、均衡分析
    add_heading(doc, "三、最优反应与均衡分析", 1)
    add_heading(doc, "3.1 三方静态博弈：由支付函数计算的唯一纳什均衡", 2)
    add_para(doc,
        "聚焦核心三方 {OpenAI, Meta, DeepSeek}，各方在 {O, C} 中同时选择。把表 2 的参数代入第二节的支付函数，"
        "即可对全部 8 个策略组合算出三方支付，无需任何人工填写。先看作为均衡的策略组合 (C, O, O)：OpenAI 闭源时独家"
        "占据高端 API 市场（n_C=1，API 收入 = 12/1 = 12.00），加闭源品牌 2.50、减成本 0.10×14.0×1.0 = 1.40，得 13.10；"
        "Meta 开源时无 API 收入，生态价值为 1.4×(6 − 1.5×1) = 6.30（此时 n_O=2），加开源品牌 2.00、减成本 0.10×30.5×0.6 = 1.83，"
        "得 6.47；DeepSeek 开源时生态价值为 1.2×4.5 = 5.40，加开源品牌 3.50、减成本 0.10×1.0×0.6 = 0.06，得 8.84。"
        "逐项分解见表 3。", line=LINE, space_after=3)
    add_table(doc,
        ["参与人（策略）", "API 收入", "生态价值", "品牌价值", "成本负担", "合计支付"],
        [["OpenAI（闭源）", "12.00", "0.00", "2.50", "1.40", "13.10"],
         ["Meta（开源）", "0.00", "6.30", "2.00", "1.83", "6.47"],
         ["DeepSeek（开源）", "0.00", "5.40", "3.50", "0.06", "8.84"]],
        widths=[1.8, 1.1, 1.1, 1.1, 1.1, 1.1], size=9,
        caption="表 3：纳什均衡 (C, O, O) 的支付逐项分解（由支付函数计算）。")
    add_para(doc,
        "对其余组合同样计算，固定 DeepSeek = O（下文将证明其开源严格占优）的 2×2 支付矩阵如表 4，三元组为 "
        "(OpenAI, Meta, DeepSeek)。", line=LINE, space_after=4)
    add_table(doc,
        ["OpenAI ＼ Meta", "Meta：O（开源）", "Meta：C（闭源）"],
        [["OpenAI：O（开源）", "(1.7, 4.4, 7.0)", "(2.4, 4.5, 8.8)"],
         ["OpenAI：C（闭源）", "(13.1, 6.5, 8.8)  ★ 纳什均衡", "(7.1, 1.0, 10.6)"]],
        widths=[2.3, 2.5, 2.5], size=9.5, zebra=False,
        caption="表 4：三方静态博弈支付矩阵（DeepSeek = O），数值由支付函数计算。对全部 8 个组合检验，得唯一纯策略纳什均衡 (C, O, O)。")
    add_para(doc, [
        ("均衡的稳定性可由迭代剔除严格劣势策略（可占优求解）严格证明，分三步。", {"size": BODY, "bold": True}),
        ("第一步，对 DeepSeek 而言，无论 OpenAI 与 Meta 如何选择，开源支付恒高于闭源（四种对手组合下分别为 "
         "7.04 对 4.20、8.84 对 2.20、8.84 对 2.20、10.64 对 1.53），故开源严格占优，剔除其闭源策略。"
         "第二步，在 DeepSeek 已开源的前提下，OpenAI 闭源对开源严格占优（Meta 开源时 13.1 对 1.7，Meta 闭源时 7.1 对 2.4），"
         "剔除其开源策略。第三步，在 OpenAI 闭源、DeepSeek 开源的前提下，Meta 选开源（6.5）优于闭源（1.0）。"
         "三步之后仅剩唯一组合 (C, O, O)，故该博弈可占优求解，纳什均衡唯一，且由于求解过程仅用到严格占优关系，"
         "也不存在混合策略纳什均衡。", {"size": BODY})], line=LINE, space_after=3)
    add_para(doc, [
        ("为何任何对 (C, O, O) 的单方偏离都不稳定？", {"size": BODY, "bold": True}),
        ("若 OpenAI 由闭源改为开源，它将放弃独占高端 API 市场的溢价，支付从 13.1 骤降至 1.7，故不会单方开源；"
         "若 Meta 由开源改为闭源，它既无 OpenAI 的先发客户、又无 DeepSeek 的成本优势，正面闭源竞争只会两败俱伤，"
         "支付从 6.5 跌至 1.0；若 DeepSeek 由开源改为闭源，其薄弱的闭源变现使支付从 8.8 跌至 2.2。换言之，"
         "「一家闭源垄断高端、其余开源争夺生态」恰好是各方都无法通过单方改变策略而获益的稳定结构——这正是 2023 年"
         "以来产业格局虽经 DeepSeek 冲击却始终维持「分层」而非「全面闭源」或「全面开源」的博弈论根源。"
         "可见，均衡结论 (C, O, O)（OpenAI 闭源、Meta 与 DeepSeek 开源）与 2026 年的真实格局完全吻合，"
         "且它是从成本、变现、生态、品牌四类原始量推导出来的，而非事先假定。", {"size": BODY})],
             line=LINE, space_after=4)
    add_figure(doc, 2, "图 2：三方支付矩阵热力图（固定 DeepSeek = O，红框★标注唯一纯策略纳什均衡）。数值由 §2.4 支付函数计算，由 01_static_game.py 生成。")

    add_heading(doc, "3.2 序贯博弈的子博弈完美均衡", 2)
    add_para(doc,
        "静态博弈假定三方同时行动，但历史上 OpenAI 先动、Meta 与 DeepSeek 后动。把同一支付（与模块 1 同源，避免二次录入）"
        "放入三阶段扩展式博弈并逆向归纳：第三阶段 DeepSeek 在任何历史下恒选开源；第二阶段 Meta 在 OpenAI 开源时选闭源"
        "（差异化反应），在 OpenAI 闭源时选开源（破局反应）；第一阶段 OpenAI 预判后选闭源（13.1 高于改为开源的 1.7）。"
        "因此子博弈完美均衡（SPE）路径为 (C, O, O)，与同时博弈的纳什均衡完全一致。两种求解的一致性说明：该均衡不依赖于"
        "「同时」这一简化假设，无论把博弈看作同时还是序贯，结论都稳健。序贯视角还额外解释了一个现实细节——OpenAI 的"
        "「先动 + 闭源」锁定了 API 基准设定者地位与开发者心智这一无形资产，故它仅以 gpt-oss 在边缘开放权重上对冲，"
        "而坚守旗舰闭源，形成「旗舰闭源 + 边缘开放」的双线策略。", line=LINE, space_after=4)
    add_figure(doc, 3, "图 3：序贯博弈扩展式博弈树（红色路径为子博弈完美均衡，逆向归纳得 (C, O, O)）。由 02_sequential_game.py 生成。")

    # ========================================================= 四、改进一
    add_heading(doc, "四、Game 改进一：参与人扩展——引入监管者（演化博弈）", 1)
    add_para(doc,
        "现实监管已实质影响开源/闭源选择：欧盟《人工智能法案》的 GPAI 义务自 2025 年 8 月 2 日适用，第 53(2) 条对开源"
        "给予部分豁免，但系统性风险模型（训练算力超过 10²⁵ FLOPs）不在豁免之列；加州 SB 53 于 2026 年 1 月 1 日生效"
        "（针对年收入超过 5 亿美元且算力超过 10²⁶ FLOPs 者）；美国第 14179 号行政令转向去管制；中国《人工智能+》行动"
        "鼓励开源、目标 2027 年重点行业渗透率达 70%。为把监管纳入博弈，引入第六个参与人——监管者 Pᴿ，将其策略效应嵌入"
        "支付：Uᵢ' = Uᵢ − γᵢ·Rᵢ(sᵢ, aᴿ)。", line=LINE, space_after=3)
    add_para(doc, [("演化博弈复制动态。", {"bold": True, "size": BODY}),
                   ("把产业看作大量厂商构成的种群，设开源策略占比为 x，其演化遵循复制动态 dx/dt = x(1−x)[π_O(x) − π_C(x)]。"
                    "与第二节一致的力学是：开源者面临生态拥挤，故 π_O = α − βx + reg（取 α=2, β=1，β 对应生态拥挤折损 τ 的"
                    "群体版本）；闭源者享受「开源者越多、闭源高端越被衬托」的溢出，故 π_C = γ + δx（取 γ=1, δ=0.5）；reg "
                    "为监管对开源净收益的影响。令 π_O = π_C 得唯一内部演化稳定策略 ", {"size": BODY}),
                   ("x* = (α − γ + reg)/(β + δ)", {"bold": True, "size": BODY, "color": PRIMARY}),
                   ("。由于 d(π_O − π_C)/dx = −(β + δ) < 0 恒成立，该内部均衡总是渐近稳定的——任何偏离都会被复制动态拉回。",
                    {"size": BODY})], line=LINE, space_after=3)
    add_table(doc,
        ["监管情景", "reg", "稳态 x*", "现实产业格局对照"],
        [["基线（无监管）", "+0.0", "0.667", "开源主导，闭源占据高端"],
         ["欧盟（GPAI 严管）", "−0.5", "0.333", "闭源主导，开源仅长尾（Mistral 已商业化转向）"],
         ["中国（鼓励开源）", "+0.3", "0.867", "开源压倒性（DeepSeek / Qwen / GLM / Kimi / Yi 全开源）"],
         ["美国（中性偏支持）", "−0.1", "0.600", "两阵营基本平分（OpenAI 闭 + Meta 开 + gpt-oss）"]],
        widths=[1.7, 0.8, 1.0, 3.7], size=9,
        caption="表 5：四种监管情景下的 ESS 稳态 x*。监管把唯一均衡分裂为四个区域性 ESS，其预测值（0.33–0.87）与实际观测高度吻合，构成模型的可证伪检验。")
    add_para(doc, [
        ("四个区域 ESS 的现实映射清晰可辨。", {"size": BODY, "bold": True}),
        ("在欧盟，GPAI 合规成本抬高了开源的边际负担，Mistral 逐步转向 API 商业化、部分厂商规避在欧盟域内发布权重，"
         "开源占比被压低（x* ≈ 0.33）；在中国，《人工智能+》行动与产业政策的鼓励下，DeepSeek、Qwen、智谱 GLM、"
         "Moonshot Kimi、零一万物 Yi 等几乎全部采用 Apache 2.0 或 MIT，形成压倒性开源生态（x* ≈ 0.87）；在美国，"
         "去管制与加州 SB 53 透明度要求相互对冲，呈现 OpenAI 闭源旗舰 + gpt-oss 开放权重 + Meta 开源 + Anthropic "
         "严格闭源的均衡分布（x* ≈ 0.60）。跨区差异还催生了「监管套利」——在监管宽松地训练、在受限地受限部署，"
         "已成为模型公司的现实选择。这里需要诚实说明：reg 取值是按各辖区监管的方向与力度标定的，并非独立测得的硬数据；"
         "模型的检验力体现在「给定方向标定后，ESS 的数值预测与各区实际开源比例一致」这一可证伪对照上。", {"size": BODY})],
             line=LINE, space_after=4)
    add_figure(doc, 4, "图 4：复制动态相图（(a) 唯一内部 ESS）与四种监管情景下开源占比的演化轨迹（(b) scipy 数值积分）。由 05_evolutionary_regulator.py 生成。")

    # ========================================================= 五、改进二
    add_heading(doc, "五、Game 改进二：信息结构——贝叶斯信号博弈", 1)
    add_para(doc,
        "披露本身是一种战略选择，构成 Spence 式信号博弈。DeepSeek 的 V3 技术报告完整公开了 MoE + MLA + FP8 训练细节、"
        "约 278.8 万 H800 GPU 小时与 557.6 万美元成本；由于全球研究者可用 MIT 权重复现，伪造这些数字的代价极高，"
        "因而是典型的可验证强信号。设强类型（真实低成本）发信号成本 c_S 较低、弱类型（伪造）成本 c_W 较高，"
        "被市场识别为强类型可获额外收益 ΔU = 4。", line=LINE, space_after=3)
    add_para(doc, [("分离均衡的充要条件为 ", {"size": BODY}),
                   ("c_S < ΔU < c_W", {"bold": True, "size": 11, "color": PRIMARY}),
                   ("：强类型愿发、弱类型不愿伪造。据此把各案例按 (c_S, c_W) 落点分为三区（见图 5 右）。"
                    "其一为分离区（信号有效）：DeepSeek (1.5, 7.5)、Llama 3.1-405B (1.0, 8.5)、Qwen3 (2.0, 9.0)，"
                    "可复现的开源使伪造门槛足够高；其二为混同且均发信号（c_W < ΔU）：营销夸大的小模型 (0.6, 1.8)，"
                    "伪造门槛过低致信号被稀释；其三为混同且均不发信号（c_S > ΔU）：无差异化跟随者 (6.0, 6.6)，"
                    "披露门槛过高致强类型也选择沉默。可见「可复现的开源」远比「单纯宣称低成本」更具说服力。",
                    {"size": BODY})], line=LINE, space_after=3)
    add_table(doc,
        ["阶段", "关键事件", "时间", "后验 P(强类型)"],
        [["T₀", "2024 年初先验", "2024-Q1", "0.100"],
         ["T₁", "V2 论文（MoE + MLA）", "2024-05", "0.270"],
         ["T₂", "V3 + 详细技术报告（557.6 万美元）", "2024-12", "0.759"],
         ["T₃", "R1 / R1-Zero 发布", "2025-01", "0.984"],
         ["T₄", "全球复现 + App Store 登顶", "2025-02", "≈1.000"]],
        widths=[0.8, 3.0, 1.2, 2.2], size=9,
        caption="表 6：DeepSeek 类型后验信念的贝叶斯演化。信念跃迁使「维持高价」不再是 OpenAI 的最优反应，正是 2025 年 API 价格雪崩与 gpt-oss 推出的信号博弈解释。")
    add_para(doc,
        "信号博弈与定价的联动机制可这样理解。在先验 µ₀ = 0.10 下，市场普遍怀疑「低成本前沿模型」的可行性，因此 OpenAI "
        "维持 GPT-4 的高价是贝叶斯均衡上的最优反应——高价既是收入来源，也是「高成本等于高质量」的隐性信号。然而当 "
        "DeepSeek 用可复现的开源权重把后验推高至接近 1 之后，「高价等于高质量」的信号被打破：买方意识到同等能力可以以"
        "极低成本获得，OpenAI 若继续维持高价只会加速客户流失。于是降价从「不可置信」变为「占优」，这解释了 GPT-5 输入价"
        "较 GPT-4 暴跌约 95.8%，以及 OpenAI 罕见地以 gpt-oss 重新进入开放权重领域以重建开源声誉。这里同样需要诚实标注："
        "先验与各阶段的似然比是按市场反应的可观测序列标定的，而贝叶斯更新本身（由先验与似然推出后验）是精确计算的。",
        line=LINE, space_after=4)
    add_figure(doc, 5, "图 5：贝叶斯信念演化（(a) 先验 0.10 → 后验 ≈1.00）与 Spence 信号三区域划分（(b)）。由 03_bayesian_signal.py 生成。")

    # ========================================================= 六、改进三
    add_heading(doc, "六、Game 改进三：行动顺序——无限期重复博弈", 1)
    add_para(doc,
        "大模型迭代是持续过程（GPT 约每 12 个月一代、DeepSeek 约每 3 至 6 个月一代），故可用无限期重复博弈分析 OpenAI "
        "与 Anthropic 在高端 API 定价上的隐性协调。与外生设定四个单期支付不同，本文先用一个差异化伯特兰定价模型把它们"
        "推导出来。设两家闭源前沿厂商面临线性需求 qᵢ = a − pᵢ + g·pⱼ，其中 g ∈ (0,1) 为产品替代度，边际成本 mc ≈ 0"
        "（前沿模型的推理边际成本相对 API 价格极小，是标准的软件经济学简化）。取 a = 12、g = 0.5、mc = 0。", line=LINE, space_after=3)
    add_para(doc, [("由该模型可解出四种情形的价格与单期利润。", {"size": BODY, "bold": True}),
                   ("合作（双方设联合利润最大价）的合谋价 pᵐ = a/[2(1−g)] = 12，单期利润 C = 72；一次性伯特兰纳什价 "
                    "p* = a/(2−g) = 8，对应价格战（惩罚）利润 P = 64；背叛者在对手维持 pᵐ 时的最优反应价 p_d = (a+g·pᵐ)/2 = 9，"
                    "背叛当期利润 D = 81；守约一方被背叛时的利润 S = 54。四者满足囚徒困境的标准排序 D > C > P > S "
                    "（81 > 72 > 64 > 54），无需人为指定。", {"size": BODY})], line=LINE, space_after=3)
    add_para(doc,
        "在冷酷触发策略下，维持合作的激励相容条件 C/(1−δ) ≥ D + δ·P/(1−δ) 化简得临界贴现因子：", line=LINE, space_after=2)
    add_para(doc, [("δ* = (D − C)/(D − P) = (81 − 72)/(81 − 64) = 9/17 ≈ 0.529",
                    {"bold": True, "size": 11, "color": PRIMARY})],
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    add_para(doc, [("双因子有效贴现。", {"bold": True, "size": BODY}),
                   ("现实中两股力量同时压低闭源阵营的有效贴现因子：外部颠覆强度 I（DeepSeek 冲击使未来份额高度不确定）"
                    "与 IPO 短视 λ（2026 年两家筹备上市，融资压力抬高当期权重）。建模为 δ_eff = δ_base·(1−I)·(1−λ)。"
                    "取 δ_base = 0.85、I = 0.30、λ = 0.30，则 δ_eff ≈ 0.416，已低于 δ* ≈ 0.529，合谋无法维持；"
                    "若暂不计 IPO（λ = 0），则单独求出临界颠覆强度 I* = 1 − δ*/δ_base ≈ 0.377——颠覆侵蚀一旦超过约 37.7%，"
                    "默契高价即告破裂。", {"size": BODY})], line=LINE, space_after=3)
    add_para(doc,
        "这一推导与现实高度吻合：DeepSeek V3（I 约 0.30）到 R1 全球复现叠加 NVIDIA 市值蒸发约 5890 亿美元（I 约 0.50，"
        "已超过 I*），再到 GPT-5 输入价下降约 95.8%，最后是 2026 年 IPO 短视叠加、价格战常态化；Stanford AI Index 2025 "
        "记录的 18 个月 280 倍推理降价，正是上述机制长期累积的结果。在声誉机制方面，截至 2026 年 5 月，阿里 Qwen 仍全面"
        "开源（约占全球开源下载的一半，衍生模型数自 2025 年 10 月起超越 Llama），头部开源者尚无旗舰回归闭源的实例，"
        "与「中国高开源 ESS（0.867）」的预测一致。右图用 2023–2026 年真实前沿旗舰输入价刻画了「阶段一价格战崩溃、"
        "阶段二前沿企稳回升、长尾持续下探」的价格分层格局。", line=LINE, space_after=4)
    add_figure(doc, 6, "图 6：双因子贴现瀑布图（(a) 跌破临界线 δ*）与前沿旗舰输入价时间线（(b) 对数轴「崩溃→分层回升」）。单期支付由差异化伯特兰模型推导，由 04_repeated_game.py 生成。")

    # ========================================================= 七、结论
    add_heading(doc, "七、结论与现实预测", 1)
    add_table(doc,
        ["改进维度", "原均衡", "改进后均衡", "关键变化"],
        [["基线（静态/序贯）", "纳什唯一", "(C, O, O)", "由支付函数推导的单一稳态"],
         ["+ 监管者（演化）", "单一", "4 个区域 ESS", "政策依赖性显现"],
         ["+ 信号博弈", "外生类型", "Spence 分离均衡", "透明度成为战略工具"],
         ["+ 重复博弈", "一次性", "δ* ≈ 0.529 临界", "由伯特兰模型量化合谋脆弱性"]],
        widths=[1.8, 1.3, 1.9, 2.2], size=9,
        caption="表 7：四类博弈模型的均衡变化对比——本文在同一产业中叠加应用了静态、序贯、贝叶斯信号、重复与演化五类模型。")
    add_para(doc, [("现实预测（2026–2027）。", {"bold": True, "size": BODY}),
                   ("第一，闭源阵营将进一步分层为「前沿闭源、追赶半开、长尾全开」，gpt-oss 模式向更多厂商扩散；"
                    "第二，监管套利推动产业地理重组，训练与部署在不同辖区分离；第三，新颠覆者（具身智能、世界模型、"
                    "多模态视频）将再次重置贴现因子 δ，若 DeepSeek 后续版本持续逼近前沿，可能触发新一轮降价；"
                    "第四，开源生态的「分叉成本」上升，Llama、Qwen、DeepSeek 各自形成社区与工具链，竞争从同质化转为"
                    "多生态平行并存；第五，估值与营收的领先权动态再分配——Anthropic 企业端领先、OpenAI 消费端领先的"
                    "双轨格局延续，两家的 IPO 进程将成为下一阶段支付函数的关键变量。", {"size": BODY})],
             line=LINE, space_after=3)
    add_para(doc, [("理论意义。", {"bold": True, "size": BODY}),
                   ("本研究的价值不在于任何单个模型，而在于展示了博弈论工具箱在同一产业问题上的「叠加解释力」：静态博弈"
                    "定位基础均衡、序贯博弈刻画先动优势、贝叶斯信号解释信息冲击、重复博弈量化合谋的脆弱性、演化博弈揭示"
                    "监管的区域分化。五类模型彼此印证，且每一步结论都能对照 2023–2026 年的真实事件加以检验，体现了博弈论"
                    "作为「可证伪的社会科学方法」的解释深度。更重要的是，本版把静态支付与重复博弈单期支付都建立在可追溯的"
                    "原始量之上，使「均衡」不再是事先假定，而是从成本、变现、生态与定价中推导出来的结果。", {"size": BODY})],
             line=LINE, space_after=3)
    add_para(doc, [("研究局限。", {"bold": True, "size": 10}),
                   ("除训练成本比为硬数据外，支付函数中的其余参数与演化博弈的 reg、信号博弈的先验与似然均为按公开证据"
                    "标定的序数量，其相对排序稳健而绝对数值仅供计算；本文未显式建模用户侧迁移博弈；地缘政治（芯片管制）"
                    "与 AI 估值的循环融资、交叉持股亦未纳入，均可作为后续扩展。尽管如此，本研究表明博弈论能以简洁模型"
                    "刻画复杂产业格局，并给出可证伪的预测。", {"size": 10})], line=LINE, space_after=3)
    add_para(doc, [
        ("总结。", {"bold": True, "size": BODY}),
        ("「开源已成为不可逆的均衡力量，闭源前沿与开源生态将长期分层并存」——这一格局可由本文的多层博弈框架统一解释，"
         "其核心均衡 (C, O, O) 由支付函数推导、其合谋脆弱性由伯特兰模型量化、其区域差异由演化博弈预测，且全部图表均可"
         "由附带代码 1:1 复现。", {"size": BODY})], line=LINE, space_after=4)

    # zoom 修正（python-docx 模板默认 <w:zoom w:val="bestFit"/>，严格 OOXML 需要 w:percent）
    try:
        zoom = doc.settings.element.find(qn("w:zoom"))
        if zoom is not None:
            for a in list(zoom.attrib):
                del zoom.attrib[a]
            zoom.set(qn("w:percent"), "100")
    except Exception as exc:  # pragma: no cover
        print("zoom fix skipped:", exc)

    doc.save(OUT_DOCX)
    print("Saved:", OUT_DOCX)


if __name__ == "__main__":
    build()
